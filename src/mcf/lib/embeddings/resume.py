"""Resume text extraction helpers."""

from __future__ import annotations

import io
from pathlib import Path


def extract_resume_text(source: str | Path | bytes) -> str:
    """Extract plain text from a resume file.

    *source* may be:
      - a file path (str or Path) — used by the local CLI flow
      - raw bytes — used by the upload endpoint (no temp file needed)

    Supported formats: .pdf, .docx, .txt, .md
    When *source* is bytes the format is detected by sniffing magic bytes
    (PDF starts with ``%PDF``; DOCX is a ZIP).
    """
    if isinstance(source, (str, Path)):
        return _extract_from_path(Path(source))
    return _extract_from_bytes(source)


def _extract_from_path(p: Path) -> str:
    suffix = p.suffix.lower()

    if suffix in {".txt", ".md"}:
        return p.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(p))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix == ".docx":
        from docx import Document  # type: ignore

        doc = Document(str(p))
        return "\n".join(par.text for par in doc.paragraphs if par.text)

    raise ValueError(f"Unsupported resume file type: {suffix} (supported: .txt, .md, .pdf, .docx)")


def _extract_from_bytes(data: bytes) -> str:
    """Detect format from magic bytes and extract text."""
    if data[:4] == b"%PDF":
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    # DOCX is a ZIP file starting with PK\x03\x04
    if data[:2] == b"PK":
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(data))
        return "\n".join(par.text for par in doc.paragraphs if par.text)

    # Fall back: treat as plain text
    return data.decode("utf-8", errors="ignore")
